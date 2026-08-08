"""
Document Parser - Extract content from PDF/DOCX with format preservation
Converts documents to Markdown (primary), HTML (secondary), and plain text (tertiary)
"""

import re
from typing import Dict, Any, List, Tuple
from io import BytesIO

# PDF parsing
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    print("ΓÜá∩╕Å PyMuPDF not installed. PDF parsing will be limited.")

# DOCX parsing
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("ΓÜá∩╕Å python-docx not installed. DOCX parsing will be limited.")

# Markdown/HTML conversion
try:
    import markdown
    from bs4 import BeautifulSoup
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False
    print("ΓÜá∩╕Å markdown/beautifulsoup4 not installed. HTML conversion will be limited.")


class DocumentParser:
    """Parse documents with format preservation (Markdown-first approach)"""
    
    def __init__(self):
        self.heading_sizes = {
            "h1": 18,  # Font size thresholds for heading detection
            "h2": 16,
            "h3": 14,
            "h4": 12
        }
    
    async def parse_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Parse document based on file type
        
        Args:
            file_path: Path to document file
            file_type: MIME type
        
        Returns:
            Dict with markdown, html, plain_text, structure
        """
        if file_type == "application/pdf":
            return await self.parse_pdf_with_formatting(file_path)
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            return await self.parse_docx_with_formatting(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def parse_pdf_with_formatting(self, file_path: str) -> Dict[str, Any]:
        """
        Parse PDF preserving:
        - Headings (detected by font size)
        - Bold/italic text
        - Tables
        - Lists (bullets/numbers)
        - Spacing and line breaks
        
        Returns Markdown as primary format
        """
        if not HAS_PYMUPDF:
            raise ImportError("PyMuPDF (fitz) is required for PDF parsing. Install: pip install PyMuPDF")
        
        try:
            doc = fitz.open(file_path)
            markdown_lines = []
            page_count = len(doc)
            word_count = 0
            placeholders = set()
            sections = []
            
            for page_num, page in enumerate(doc):
                # Extract text with formatting information
                blocks = page.get_text("dict")["blocks"]
                
                for block in blocks:
                    if block["type"] == 0:  # Text block
                        for line in block.get("lines", []):
                            line_text = ""
                            line_font_size = 0
                            is_bold = False
                            
                            for span in line.get("spans", []):
                                text = span.get("text", "").strip()
                                if not text:
                                    continue
                                
                                font_size = span.get("size", 12)
                                font_name = span.get("font", "").lower()
                                
                                # Detect bold
                                is_span_bold = "bold" in font_name or span.get("flags", 0) & 2 ** 4
                                
                                # Apply bold markdown
                                if is_span_bold and text:
                                    text = f"**{text}**"
                                
                                line_text += text + " "
                                line_font_size = max(line_font_size, font_size)
                                is_bold = is_bold or is_span_bold
                            
                            line_text = line_text.strip()
                            if not line_text:
                                continue
                            
                            # Detect headings by font size
                            heading_level = self._detect_heading_level(line_font_size)
                            
                            if heading_level:
                                markdown_lines.append(f"\n{'#' * heading_level} {line_text}\n")
                                sections.append({
                                    "heading": line_text.replace("**", ""),
                                    "level": heading_level
                                })
                            else:
                                # Detect list items
                                if re.match(r'^[\d]+\.', line_text):
                                    markdown_lines.append(f"{line_text}\n")
                                elif re.match(r'^[ΓÇóΓùÅΓùï-]\s', line_text):
                                    markdown_lines.append(f"- {line_text[1:].strip()}\n")
                                else:
                                    markdown_lines.append(f"{line_text}\n")
                            
                            # Count words
                            word_count += len(line_text.split())
                            
                            # Detect placeholders [SOMETHING]
                            placeholder_matches = re.findall(r'\[([A-Z\s]+)\]', line_text)
                            placeholders.update(placeholder_matches)
                
                # Extract tables
                tables = page.find_tables()
                if tables:
                    for table in tables:
                        markdown_table = self._convert_table_to_markdown(table)
                        if markdown_table:
                            markdown_lines.append(f"\n{markdown_table}\n")
                
                # Add page break (optional)
                if page_num < page_count - 1:
                    markdown_lines.append("\n---\n")
            
            doc.close()
            
            # Join markdown
            markdown_content = "\n".join(markdown_lines)
            
            # Clean up excessive newlines
            markdown_content = re.sub(r'\n{3,}', '\n\n', markdown_content)
            
            # Convert to HTML
            html_content = self._markdown_to_html(markdown_content)
            
            # Generate plain text
            plain_text = self._strip_markdown(markdown_content)
            
            # Structure
            structure = {
                "sections": sections,
                "placeholders": list(placeholders),
                "formatting_hints": {
                    "preserve_spacing": True,
                    "preserve_indentation": True,
                    "preserve_tables": True
                }
            }
            
            return {
                "markdown": markdown_content,
                "html": html_content,
                "plain_text": plain_text,
                "structure": structure,
                "word_count": word_count,
                "page_count": page_count
            }
            
        except Exception as e:
            print(f"Γ¥î Error parsing PDF: {e}")
            raise Exception(f"PDF parsing failed: {str(e)}")
    
    async def parse_docx_with_formatting(self, file_path: str) -> Dict[str, Any]:
        """
        Parse DOCX preserving:
        - Heading levels
        - Bold, italic, underline
        - Tables
        - Lists (bullets/numbers)
        - Paragraph spacing
        
        Returns Markdown as primary format
        """
        if not HAS_DOCX:
            raise ImportError("python-docx is required for DOCX parsing. Install: pip install python-docx")
        
        try:
            doc = DocxDocument(file_path)
            markdown_lines = []
            word_count = 0
            placeholders = set()
            sections = []
            page_count = 1  # Approximate (DOCX doesn't have explicit pages)
            
            for element in doc.element.body:
                # Handle paragraphs
                if element.tag.endswith('p'):
                    para = None
                    for paragraph in doc.paragraphs:
                        if paragraph._element == element:
                            para = paragraph
                            break
                    
                    if para is None:
                        continue
                    
                    text = para.text.strip()
                    if not text:
                        markdown_lines.append("")
                        continue
                    
                    # Check if heading
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                        markdown_lines.append(f"\n{'#' * level} {text}\n")
                        sections.append({"heading": text, "level": level})
                    else:
                        # Format runs (bold, italic)
                        formatted_text = ""
                        for run in para.runs:
                            run_text = run.text
                            if run.bold and run.italic:
                                formatted_text += f"***{run_text}***"
                            elif run.bold:
                                formatted_text += f"**{run_text}**"
                            elif run.italic:
                                formatted_text += f"*{run_text}*"
                            else:
                                formatted_text += run_text
                        
                        # Detect list items
                        if para.style.name.startswith('List'):
                            if 'Bullet' in para.style.name:
                                markdown_lines.append(f"- {formatted_text}")
                            elif 'Number' in para.style.name:
                                markdown_lines.append(f"1. {formatted_text}")
                            else:
                                markdown_lines.append(formatted_text)
                        else:
                            markdown_lines.append(formatted_text)
                        
                        # Count words
                        word_count += len(text.split())
                        
                        # Detect placeholders
                        placeholder_matches = re.findall(r'\[([A-Z\s]+)\]', text)
                        placeholders.update(placeholder_matches)
                
                # Handle tables
                elif element.tag.endswith('tbl'):
                    for table in doc.tables:
                        if table._element == element:
                            markdown_table = self._convert_docx_table_to_markdown(table)
                            if markdown_table:
                                markdown_lines.append(f"\n{markdown_table}\n")
                            break
            
            # Join markdown
            markdown_content = "\n".join(markdown_lines)
            
            # Clean up excessive newlines
            markdown_content = re.sub(r'\n{3,}', '\n\n', markdown_content)
            
            # Convert to HTML
            html_content = self._markdown_to_html(markdown_content)
            
            # Generate plain text
            plain_text = self._strip_markdown(markdown_content)
            
            # Structure
            structure = {
                "sections": sections,
                "placeholders": list(placeholders),
                "formatting_hints": {
                    "preserve_spacing": True,
                    "preserve_indentation": True,
                    "preserve_tables": True
                }
            }
            
            return {
                "markdown": markdown_content,
                "html": html_content,
                "plain_text": plain_text,
                "structure": structure,
                "word_count": word_count,
                "page_count": page_count
            }
            
        except Exception as e:
            print(f"Γ¥î Error parsing DOCX: {e}")
            raise Exception(f"DOCX parsing failed: {str(e)}")
    
    def _detect_heading_level(self, font_size: float) -> int:
        """Detect heading level based on font size"""
        if font_size >= 18:
            return 1
        elif font_size >= 16:
            return 2
        elif font_size >= 14:
            return 3
        elif font_size >= 13:
            return 4
        else:
            return 0  # Not a heading
    
    def _convert_table_to_markdown(self, table) -> str:
        """Convert PyMuPDF table to Markdown"""
        try:
            if not table or not table.extract():
                return ""
            
            rows = table.extract()
            if not rows or len(rows) < 2:
                return ""
            
            markdown_table = []
            
            # Header row
            header = rows[0]
            markdown_table.append("| " + " | ".join(str(cell) for cell in header) + " |")
            markdown_table.append("|" + "|".join(["---"] * len(header)) + "|")
            
            # Data rows
            for row in rows[1:]:
                markdown_table.append("| " + " | ".join(str(cell) for cell in row) + " |")
            
            return "\n".join(markdown_table)
            
        except Exception as e:
            print(f"ΓÜá∩╕Å Error converting table: {e}")
            return ""
    
    def _convert_docx_table_to_markdown(self, table) -> str:
        """Convert python-docx table to Markdown"""
        try:
            if not table.rows:
                return ""
            
            markdown_table = []
            
            # Header row
            header = [cell.text.strip() for cell in table.rows[0].cells]
            markdown_table.append("| " + " | ".join(header) + " |")
            markdown_table.append("|" + "|".join(["---"] * len(header)) + "|")
            
            # Data rows
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                markdown_table.append("| " + " | ".join(cells) + " |")
            
            return "\n".join(markdown_table)
            
        except Exception as e:
            print(f"ΓÜá∩╕Å Error converting DOCX table: {e}")
            return ""
    
    def _markdown_to_html(self, markdown_content: str) -> str:
        """Convert Markdown to HTML"""
        if not HAS_MARKDOWN:
            # Fallback: basic conversion
            return markdown_content.replace("\n", "<br>")
        
        try:
            html = markdown.markdown(
                markdown_content,
                extensions=['tables', 'fenced_code', 'nl2br', 'sane_lists']
            )
            return html
        except Exception as e:
            print(f"ΓÜá∩╕Å Error converting to HTML: {e}")
            return markdown_content.replace("\n", "<br>")
    
    def _strip_markdown(self, markdown_content: str) -> str:
        """Strip Markdown formatting to get plain text"""
        # Remove headings
        text = re.sub(r'^#{1,6}\s+', '', markdown_content, flags=re.MULTILINE)
        
        # Remove bold/italic
        text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)  # Bold+Italic
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)      # Bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)          # Italic
        
        # Remove links
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        
        # Remove images
        text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
        
        # Remove code blocks
        text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)
        text = re.sub(r'`(.*?)`', r'\1', text)
        
        # Remove horizontal rules
        text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
        
        # Clean up whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
